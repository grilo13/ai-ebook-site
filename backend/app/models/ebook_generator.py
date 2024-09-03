import json
import os
import subprocess
import uuid

import pypdf
import pypandoc
from docx import Document
from docx.shared import Mm
from docxtpl import DocxTemplate, InlineImage

from backend.app.models.ebook import Ebook
from backend.app.models.langchain_wrapper import LangchainWrapper
from backend.app.models.pdf_converter import PDFConverter


class EbookGenerator:
    def __init__(self, id: str, output_directory: str):
        self.langchain_wrapper = LangchainWrapper()
        self.pdf_converter = PDFConverter()
        self.output_directory = output_directory
        self.cover_photo_location = (
            f"cover_photo-{id}.jpg"
        )

        self.content_pdf_location = "backend/app/temp/"
        self.cover_pdf_location = "backend/app/temp/"
        self.final_pdf_location = f"backend/app/temp/book-{id}.pdf"

        # cover location
        self.cover_location = f"backend/app/temp/cover-{id}.docx"
        self.cover_final_pdf_location = f"backend/app/temp/cover-{id}.pdf"

        # final pdf
        self.final_pdf_output = f"backend/app/output/final-{id}.pdf"
        self.final_pdf_preview = f"backend/app/preview/final-{id}.pdf"

    def generate_title(self, topic: str, target_audience: str) -> str:
        print("starting generate title")
        title_prompt = (
            f'We are writing an eBook. It is about "{topic}". Our'
            f' reader is:  "{target_audience}". Write a short, catch'
            " title clearly directed at our reader that is less than"
            " 9 words and proposes a “big promise” that will be sure to grab"
            " the readers attention."
        )
        title = self.langchain_wrapper.generate_completion(prompt=title_prompt)
        # remove surrounding quotes from title (if any)
        title = title.replace('"', "")
        return title

    def generate_cover_photo(self,
                             title: str,
                             topic: str,
                             target_audience: str,
                             img_output: str):
        cover_prompt = (
            f'We have a ebook with the title {title}. It is about "{topic}".'
            f' Our reader is:  "{target_audience}". Write me a very brief and'
            " matter-of-fact description of a photo that would be on the"
            " cover of the book. Do not reference the cover or photo in your"
            ' answer. For example, if the title was "How to lose weight for'
            ' middle aged women", a reasonable response would be "a middle'
            ' age woman exercising"'
        )
        print("cover prompt", cover_prompt)
        photo_prompt = self.langchain_wrapper.generate_completion(prompt=cover_prompt)
        print("dalle prompt", photo_prompt)
        img_data = self.langchain_wrapper.generate_photo(photo_prompt=photo_prompt)
        with open(img_output, "wb") as f:
            f.write(img_data)

    def generate_cover(self, cover_template: str,
                       title: str,
                       topic: str,
                       target_audience: str,
                       output_file: str,
                       preview: bool):
        doc = DocxTemplate(cover_template)

        if preview:
            self.cover_photo_location = "backend/app/templates/covers/preview_photo.png"
        else:
            self.generate_cover_photo(title=title, topic=topic, target_audience=target_audience,
                                      img_output=self.cover_photo_location)

        imagen = InlineImage(
            doc, self.cover_photo_location, width=Mm(120)
        )  # width is in millimetres
        context = {"title": title, "subtext": "Ebook Generator", "image": imagen}
        doc.render(context)
        doc.save(self.cover_location)
        self.pdf_converter.convert_to(docx=self.cover_location, folder=output_file)

    def generate_outline(
            self,
            topic,
            target_audience,
            title,
            num_chapters,
            num_subsections,
    ):
        outline_prompt = (
            f'We are writing an eBook called "{title}". It is about'
            f' "{topic}". Our reader is:  "{target_audience}".  Create'
            " a compehensive outline for our ebook, which will have"
            f" {num_chapters} chapter(s). Each chapter should have exactly"
            f" {num_subsections} subsection(s) Output Format for prompt:"
            " python dict with key: chapter title, value: a single list/array"
            " containing subsection titles within the chapter (the subtopics"
            " should be inside the list). The chapter titles should be"
            ' prepended with the chapter number, like this: "Chapter 5:'
            ' [chapter title]". The subsection titles should be prepended'
            ' with the {chapter number.subtitle number}, like this: "5.4:'
            ' [subsection title]". '
        )
        outline_json = self.langchain_wrapper.generate_completion(prompt=outline_prompt)
        outline_json = outline_json[outline_json.find("{"): outline_json.rfind("}") + 1]
        outline = json.loads(outline_json)
        print("outline", outline)
        """if not self.verify_outline(outline, num_chapters, num_subsections):
            raise Exception("Outline not well formed!")"""

        return outline

    def generate_chapter_content(self, topic: str,
                                 target_audience: str,
                                 title: str,
                                 idx: int,
                                 chapter: str,
                                 subtopic: str):
        num_words_str = "500 to 700"
        content_prompt = (
            f'We are writing an eBook called "{title}". Overall, it is about'
            f' "{topic}". Our reader is:  "{target_audience}". We are'
            f" currently writing the #{idx + 1} section for the chapter:"
            f' "{chapter}". Using at least {num_words_str} words, write the'
            " full contents of the section regarding this subtopic:"
            f' "{subtopic}". The output should be as helpful to the reader as'
            " possible. Include quantitative facts and statistics, with"
            " references. Go as in depth as necessary. You can split this"
            " into multiple paragraphs if you see fit. The output should also"
            ' be in cohesive paragraph form. Do not include any "[Insert'
            ' ___]" parts that will require manual editing in the book later.'
            " If find yourself needing to put 'insert [blank]' anywhere, do"
            " not do it (this is very important). If you do not know"
            " something, do not include it in the output. Exclude any"
            " auxiliary information like  the word count, as the entire"
            " output will go directly into the ebook for readers, without any"
            " human processing. Remember the {num_words_str} word minimum,"
            " please adhere to it."
        )
        content = self.langchain_wrapper.generate_completion(prompt=content_prompt)
        return content

    def generate_docx(self,
                      topic: str,
                      target_audience: str,
                      title: str,
                      outline: dict,
                      docx_file: str,
                      book_template: str,
                      preview: bool,
                      actionable_steps: bool = False):
        print("generating docx")
        document = Document(book_template)
        document.add_page_break()
        document.add_heading("Table of Contents")
        for chapter, subtopics in outline.items():
            print("chapter", chapter)
            print("subtopics", subtopics)
            document.add_heading(chapter, level=2)
            for idx, subtopic in enumerate(subtopics):
                print("idx", idx)
                print("subtopic", subtopic)
                document.add_heading("\t" + subtopic, level=3)

        document.add_page_break()

        chapter_num = 1
        for chapter, subtopics in outline.items():
            document.add_heading(chapter, level=1)
            subtopics_content = []

            # Generate each subtopic content
            for idx, subtopic in enumerate(subtopics):
                # Stop writing the ebook after four subsections if preview
                print("preview and idx", preview, idx)
                if preview and idx >= 2:
                    break
                document.add_heading(subtopic, level=2)
                content = self.generate_chapter_content(
                    topic=topic, target_audience=target_audience, title=title,
                    idx=idx, chapter=chapter, subtopic=subtopic
                )
                document.add_paragraph(content)

            if preview:
                document.add_heading(
                    "Preview Completed - Purchase Full Book To Read More!"
                )
                break

            if chapter_num < len(outline.items()):
                document.add_page_break()
            chapter_num += 1

        document.add_page_break()

        document.save(docx_file)

    @staticmethod
    def remove_first_page(source_pdf: str, output_pdf: str):
        # Open the source PDF file
        with open(source_pdf, "rb") as file:
            reader = pypdf.PdfReader(file)

            # Create a PDF writer object
            writer = pypdf.PdfWriter()

            # Number of pages in the source PDF
            num_pages = reader.get_num_pages()

            # Check if the PDF has more than one page
            if num_pages < 2:
                raise ValueError(
                    "The PDF has only one page and cannot be processed."
                )

            # Loop through all pages except the first one
            for page_num in range(1, num_pages):
                # Get the page
                page = reader.get_page(page_number=page_num)

                # Add it to the writer object
                writer.add_page(page=page)

            # Write out the new PDF
            with open(output_pdf, "wb") as output_file:
                writer.write(output_file)

    def merge_pdfs(self, input_files: list, output_file: str):
        merger = pypdf.PdfMerger()

        for pdf in input_files:
            merger.append(pdf)

        merger.write(output_file)
        merger.close()

    def generate_ebook(self,
                       topic: str,
                       target_audience: str,
                       id: str,
                       num_chapters: int = 6,
                       num_subsections: int = 4,
                       preview: bool = True) -> Ebook:

        docx_file = f"backend/app/docs/docs-{id}.docx"
        title = self.generate_title(topic=topic, target_audience=target_audience)

        template = {
            "cover_template": (
                "backend/app/templates/covers/gen.docx"
            ),
            "book_template": (
                "backend/app/templates/content/theme.docx"
            ),
        }

        self.generate_cover(cover_template=template.get("cover_template"),
                            title=title,
                            topic=topic,
                            target_audience=target_audience,
                            output_file=self.cover_pdf_location,
                            preview=preview)

        outline = self.generate_outline(
            topic,
            target_audience,
            title,
            num_chapters,
            num_subsections,
        )

        self.generate_docx(topic=topic,
                           target_audience=target_audience,
                           title=title,
                           outline=outline,
                           docx_file=docx_file,
                           book_template=template.get("book_template"),
                           preview=preview)

        # convert to pdf and remove the first empty page
        pdf_location = self.pdf_converter.convert_to(docx=docx_file, folder=self.content_pdf_location)

        self.remove_first_page(source_pdf=str(os.path.join(pdf_location)), output_pdf=self.final_pdf_location)

        if 'preview' in self.output_directory:
            final_pdf_output = self.final_pdf_preview
        else:
            final_pdf_output = self.final_pdf_output

        self.merge_pdfs(input_files=[self.cover_final_pdf_location, self.final_pdf_location],
                        output_file=final_pdf_output)

        # TODO need to add number of pages to send also on the email
        return Ebook(title=title,
                     topic=topic,
                     target_audience=target_audience,
                     docx_file=docx_file,
                     pdf_file=final_pdf_output)


if __name__ == '__main__':
    ebook_generator = EbookGenerator(id=str(uuid.uuid1()),
                                     output_directory="../preview/")
    ebook = ebook_generator.generate_ebook(topic='how to lose weight',
                                           target_audience='mid age moms',
                                           id=str(uuid.uuid1()),
                                           preview=True)
    print("ebook title: ", ebook.title)
    print("ebook topic: ", ebook.topic)
